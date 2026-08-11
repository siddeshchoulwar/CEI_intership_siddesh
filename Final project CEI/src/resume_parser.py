"""
Resume Parser
-------------
Extracts raw text from PDF / DOCX / TXT resumes and pulls out
structured signals used downstream by the scoring and feedback
modules: skills, years of experience, education level, email/phone.
"""
import re
import io
import os

# Master skill vocabulary the parser looks for. Kept flat + lowercase
# so matching is a simple substring/word-boundary check.
SKILL_VOCAB = [
    "python", "java", "c++", "c#", "javascript", "typescript", "r",
    "sql", "nosql", "mongodb", "mysql", "postgresql",
    "machine learning", "deep learning", "nlp", "computer vision",
    "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "pytorch",
    "keras", "opencv", "statistics", "linear algebra", "probability",
    "data structures", "algorithms", "rest api", "flask", "fastapi",
    "django", "react", "node.js", "docker", "kubernetes", "aws", "azure",
    "gcp", "git", "github", "tableau", "power bi", "excel", "spark",
    "hadoop", "airflow", "mlops", "microservices", "system design",
    "html", "css", "bootstrap", "streamlit", "matplotlib", "seaborn",
]

EDUCATION_KEYWORDS = [
    "b.e.", "be ", "b.tech", "btech", "m.tech", "mtech", "b.sc", "bsc",
    "m.sc", "msc", "bca", "mca", "mba", "phd", "bachelor", "master",
    "diploma",
]


def extract_text(file_path: str) -> str:
    """Extract raw text from a resume file (.pdf, .docx, .txt)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        with open(file_path, "r", errors="ignore") as f:
            return f.read()


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Same as extract_text but works from an in-memory upload (Streamlit)."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        import pdfplumber
        text = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)
    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return file_bytes.decode(errors="ignore")


def _extract_pdf(file_path: str) -> str:
    import pdfplumber
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
    return "\n".join(text)


def _extract_docx(file_path: str) -> str:
    import docx
    doc = docx.Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_email(text: str) -> str:
    m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    m = re.search(r"(\+?\d{1,3}[-.\s]?)?\d{10}", text)
    return m.group(0) if m else ""


def extract_skills(text: str) -> list:
    text_l = text.lower()
    found = []
    for skill in SKILL_VOCAB:
        pattern = r"(?<![a-zA-Z])" + re.escape(skill) + r"(?![a-zA-Z])"
        if re.search(pattern, text_l):
            found.append(skill)
    # normalize sklearn -> scikit-learn to avoid double counting
    if "sklearn" in found and "scikit-learn" not in found:
        found.append("scikit-learn")
    return sorted(set(found))


def extract_experience_years(text: str) -> float:
    """
    Looks for explicit patterns like '2 years of experience' or
    '3+ years'. Falls back to 0 if nothing found (e.g. fresher resume).
    """
    text_l = text.lower()
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*years?\s+(?:of\s+)?experience",
        r"experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*years?",
    ]
    years = []
    for p in patterns:
        for m in re.finditer(p, text_l):
            try:
                years.append(float(m.group(1)))
            except ValueError:
                pass
    return max(years) if years else 0.0


def extract_education(text: str) -> list:
    text_l = text.lower()
    found = [kw.strip() for kw in EDUCATION_KEYWORDS if kw in text_l]
    return sorted(set(found))


def parse_resume(text: str) -> dict:
    """Run the full extraction pipeline on raw resume text."""
    return {
        "raw_text": text,
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "experience_years": extract_experience_years(text),
        "education": extract_education(text),
    }
